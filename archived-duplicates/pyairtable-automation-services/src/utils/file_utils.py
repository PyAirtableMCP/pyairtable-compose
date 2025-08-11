import os
import csv
import json
from typing import List, Tuple, Optional, Dict, Any
import logging

logger = logging.getLogger(__name__)

def validate_file_extension(filename: str, allowed_extensions: List[str]) -> bool:
    """Validate file extension against allowed extensions"""
    if not filename or '.' not in filename:
        return False
    
    file_extension = filename.split('.')[-1].lower()
    return file_extension in [ext.lower() for ext in allowed_extensions]

def extract_file_content(file_path: str, mime_type: str) -> Tuple[Optional[str], Optional[Dict[str, Any]]]:
    """Extract content from file based on MIME type"""
    try:
        content = None
        metadata = {
            "file_path": file_path,
            "mime_type": mime_type,
            "extraction_method": None
        }
        
        if mime_type.startswith('text/') or mime_type == 'application/csv':
            content, extraction_meta = _extract_text_content(file_path, mime_type)
            metadata.update(extraction_meta)
            
        elif mime_type == 'application/pdf':
            content, extraction_meta = _extract_pdf_content(file_path)
            metadata.update(extraction_meta)
            
        elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                          'application/msword']:
            content, extraction_meta = _extract_word_content(file_path)
            metadata.update(extraction_meta)
            
        elif mime_type in ['application/vnd.ms-excel', 
                          'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
            content, extraction_meta = _extract_excel_content(file_path)
            metadata.update(extraction_meta)
            
        else:
            logger.warning(f"Unsupported MIME type for content extraction: {mime_type}")
            metadata["extraction_method"] = "unsupported"
            
        return content, metadata
        
    except Exception as e:
        logger.error(f"Failed to extract content from {file_path}: {str(e)}")
        return None, {"error": str(e)}

def _extract_text_content(file_path: str, mime_type: str) -> Tuple[Optional[str], Dict[str, Any]]:
    """Extract content from text files"""
    metadata = {"extraction_method": "text"}
    
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                
                metadata["encoding"] = encoding
                metadata["character_count"] = len(content)
                metadata["line_count"] = content.count('\n') + 1
                
                # Special handling for CSV files
                if mime_type == 'application/csv' or file_path.endswith('.csv'):
                    csv_metadata = _analyze_csv_content(file_path, encoding)
                    metadata.update(csv_metadata)
                
                return content, metadata
                
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, try binary mode and decode with errors='ignore'
        with open(file_path, 'rb') as f:
            binary_content = f.read()
            content = binary_content.decode('utf-8', errors='ignore')
        
        metadata["encoding"] = "utf-8-with-errors"
        metadata["character_count"] = len(content)
        
        return content, metadata
        
    except Exception as e:
        logger.error(f"Failed to extract text content: {str(e)}")
        return None, {"error": str(e)}

def _analyze_csv_content(file_path: str, encoding: str) -> Dict[str, Any]:
    """Analyze CSV file structure"""
    try:
        with open(file_path, 'r', encoding=encoding) as f:
            # Detect delimiter
            sample = f.read(1024)
            sniffer = csv.Sniffer()
            delimiter = sniffer.sniff(sample).delimiter
            
            # Reset file pointer
            f.seek(0)
            
            # Read CSV
            reader = csv.reader(f, delimiter=delimiter)
            rows = list(reader)
            
            headers = rows[0] if rows else []
            data_rows = rows[1:] if len(rows) > 1 else []
            
            return {
                "csv_delimiter": delimiter,
                "csv_headers": headers,
                "csv_column_count": len(headers),
                "csv_row_count": len(data_rows),
                "csv_total_rows": len(rows)
            }
            
    except Exception as e:
        logger.error(f"Failed to analyze CSV content: {str(e)}")
        return {"csv_error": str(e)}

def _extract_pdf_content(file_path: str) -> Tuple[Optional[str], Dict[str, Any]]:
    """Extract content from PDF files"""
    metadata = {"extraction_method": "pdf"}
    
    try:
        # Try using PyPDF2 if available
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                text_content = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    text_content += page.extract_text()
                
                metadata.update({
                    "pdf_pages": len(pdf_reader.pages),
                    "pdf_library": "PyPDF2",
                    "character_count": len(text_content)
                })
                
                return text_content, metadata
                
        except ImportError:
            # Try using pdfplumber if available
            try:
                import pdfplumber
                
                text_content = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text
                
                metadata.update({
                    "pdf_pages": len(pdf.pages),
                    "pdf_library": "pdfplumber",
                    "character_count": len(text_content)
                })
                
                return text_content, metadata
                
            except ImportError:
                logger.warning("No PDF extraction library available (PyPDF2 or pdfplumber)")
                metadata["error"] = "No PDF extraction library available"
                return None, metadata
                
    except Exception as e:
        logger.error(f"Failed to extract PDF content: {str(e)}")
        return None, {"error": str(e)}

def _extract_word_content(file_path: str) -> Tuple[Optional[str], Dict[str, Any]]:
    """Extract content from Word documents"""
    metadata = {"extraction_method": "word"}
    
    try:
        # Try using python-docx if available
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            content = '\n'.join(paragraphs)
            
            # Extract tables
            table_data = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_cells = [cell.text for cell in row.cells]
                    table_rows.append(row_cells)
                table_data.append(table_rows)
            
            metadata.update({
                "word_paragraphs": len(paragraphs),
                "word_tables": len(table_data),
                "word_library": "python-docx",
                "character_count": len(content)
            })
            
            # Add table content to main content
            if table_data:
                content += "\n\n=== TABLES ===\n"
                for i, table in enumerate(table_data):
                    content += f"\nTable {i+1}:\n"
                    for row in table:
                        content += " | ".join(row) + "\n"
            
            return content, metadata
            
        except ImportError:
            logger.warning("python-docx not available for Word document extraction")
            metadata["error"] = "python-docx library not available"
            return None, metadata
            
    except Exception as e:
        logger.error(f"Failed to extract Word content: {str(e)}")
        return None, {"error": str(e)}

def _extract_excel_content(file_path: str) -> Tuple[Optional[str], Dict[str, Any]]:
    """Extract content from Excel files"""
    metadata = {"extraction_method": "excel"}
    
    try:
        # Try using openpyxl or xlrd
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            
            content = ""
            total_rows = 0
            
            for sheet_name in sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                content += f"\n=== Sheet: {sheet_name} ===\n"
                content += df.to_string(index=False)
                content += "\n"
                
                total_rows += len(df)
            
            metadata.update({
                "excel_sheets": len(sheet_names),
                "excel_sheet_names": sheet_names,
                "excel_total_rows": total_rows,
                "excel_library": "pandas",
                "character_count": len(content)
            })
            
            return content, metadata
            
        except ImportError:
            logger.warning("pandas not available for Excel extraction")
            metadata["error"] = "pandas library not available"
            return None, metadata
            
    except Exception as e:
        logger.error(f"Failed to extract Excel content: {str(e)}")
        return None, {"error": str(e)}

def cleanup_temp_files(file_paths: List[str]):
    """Clean up temporary files"""
    for file_path in file_paths:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Cleaned up temp file: {file_path}")
        except Exception as e:
            logger.error(f"Failed to clean up temp file {file_path}: {str(e)}")